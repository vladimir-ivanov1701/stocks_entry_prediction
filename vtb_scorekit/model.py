# -*- coding: utf-8 -*-

from .data import DataSamples
from .woe import WOE
from .bankmetrics import *
from ._utils import fig_to_excel, adjust_cell_width, add_suffix, rem_suffix, is_cross_name, cross_split, add_ds_folder, make_header
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import sklearn
from sklearn.model_selection import cross_val_score, StratifiedKFold
from sklearn.linear_model import LogisticRegression, SGDClassifier
from sklearn.metrics import roc_curve, auc, make_scorer
import warnings
import os
from scipy.optimize import minimize
import scipy.stats as sts
import copy
import json
from concurrent import futures
import gc
import cloudpickle
import base64
import re
import datetime
import io
import pkg_resources
try:
    import optbinning
except:
    pass

warnings.simplefilter('ignore')
plt.rc('font', family='Verdana', size=11)
try:
    plt.style.use([s for s in plt.style.available if 'darkgrid' in s][0])
except:
    pass
pd.set_option('display.precision', 3)
gc.enable()


class LogisticRegressionModel:
    """
    Классификатор лог регрессии
    """

    def __init__(self, clf=None, ds=None, transformer=None, name=None):
        """
        :param clf: классификатор модели (должен иметь метод fit() и атрибуты coef_, intercept_). При None выбирается SGDClassifier(alpha=0.001, loss='log', max_iter=100)
        :param ds: Привязанный к модели ДатаСэмпл. Если задан, то он по умолчанию будет использоваться во всех методах
        :param transformer: объект класса WOE для трансформации факторов
        :param name: название модели
        """
        self.ds = ds.copy() if isinstance(ds, DataSamples) else None
        self.transformer = transformer
        self.round_digits = 3 if transformer is None else transformer.round_woe
        self.name = name
        self.coefs = {}
        self.intercept = None
        self.features = []
        self.calibration = None
        self.scale = None
        self.clf = clf if clf is not None else SGDClassifier(loss='log' if sklearn.__version__ < '1.' else 'log_loss',
                                                             penalty='l2', max_iter=1000, alpha=0.001,
                                                             random_state=self.ds.random_state if self.ds is not None else 0)
        if self.transformer is not None and self.ds is not None:
            self.transformer.transform(self.ds, verbose=True)
        self.print_log(f'Chosen model classifier is {self.clf}')

    def save_model(self, file_name='model.json', pickle_protocol=4):
        """
        Сохранение факторов, коэффициентов, калибровки, шкалы и биннинга в файл
        :param file_name: название json файла для сохранения модели. При None json возвращается методом
        :param pickle_protocol: версия протокола для сериализации объектов. Версия 5 доступна только для питона 3.8+
        """
        model = {p: self.__dict__[p] for p in ['name', 'coefs', 'intercept', 'calibration', 'scale']}
        model['clf'] = base64.b64encode(cloudpickle.dumps(self.clf, protocol=pickle_protocol)).decode()
        if self.transformer is not None and self.coefs:
            model['scorecard'] = base64.b64encode(cloudpickle.dumps(self.transformer.export_scorecard(features=[rem_suffix(f) for f in self.coefs], full=False).drop(['target_rate', 'sample_part'], axis=1), protocol=pickle_protocol)).decode()
        if file_name is not None:
            file_name = add_ds_folder(self.ds, file_name)
            with open(file_name, 'w', encoding='utf-8') as file:
                json.dump(model, file, ensure_ascii=False, indent=4, cls=NpEncoder)
            self.print_log(f'The model was successfully saved to file {file_name}')
        else:
            return model

    def load_model(self, file_name='model.json', verbose=True):
        """
        Загрузка факторов, коэффициентов, калибровки, шкалы и биннинга из файла
        :param file_name: название json файла или словарь для загрузки модели
        """
        if isinstance(file_name, str):
            model = json.load(open(file_name, 'rt', encoding='utf-8'))
            self.print_log(f'The model was loaded from file {file_name}')
        elif isinstance(file_name, dict):
            model = file_name
        else:
            self.print_log('file_name type must be str or dict!', 40)
            return None
        if 'clf' in model:
            self.clf = cloudpickle.loads(base64.b64decode(model['clf']))
        else:
            self.clf = LogisticRegression(random_state=model['random_state'] if 'random_state' in model else 0,
                                          C=model['regularization_value'] if 'regularization_value' in model else 1000,
                                          solver=model['solver'] if 'solver' in model else 'saga',
                                          penalty=model['regularization'] if 'regularization' in model else 'l2')
        if verbose:
            self.print_log(f'clf = {self.clf}')
        for p in ['name', 'coefs', 'intercept', 'calibration', 'scale']:
            if p in model:
                self.__dict__[p] = model[p]
                if verbose:
                    self.print_log(f'{p} = {model[p]}')
        if 'scorecard' in model:
            try:
                try:
                    scorecard = pd.DataFrame(model['scorecard'])
                except:
                    scorecard = cloudpickle.loads(base64.b64decode(model['scorecard']))
                self.transformer = WOE(ds=self.ds, scorecard=scorecard)
                self.round_digits = self.transformer.round_woe
            except Exception as e:
                self.print_log(e, 40)
                self.print_log('Error while loading scorecard! Set transformer=None', 30)
                self.transformer = None
        else:
            self.round_digits = 3
        if self.coefs:
            self.features = list(self.coefs)

    def auto_logreg(self, data=None, target=None, time_column=None, id_column=None, feature_descriptions=None, n_jobs=None,
                    result_folder='', method='opt', validate=False, out='auto_model.xlsx', save_model='auto_model.json',
                    config=None):
        """
        Построение модели в автоматическом режиме с минимальным набором параметров
        --- Выборка ---
        :param data: ДатаФрейм или ДатаСэмпл.
                - если передается ДатаФрейм, то он разбивается на трейн/тест 70%/30%
                - если передается ДатаСэмпл, то все остальные параметры блока "Выборка" не используются
        :param target: целевая переменная
        :param time_column: дата среза
        :param id_column: уникальный в рамках среза айди наблюдения
        :param feature_descriptions: датафйрем с описанием переменных. Должен содержать индекс с названием переменных и любое кол-во полей с описанием, которые будут подтягиваться в отчеты

        --- Параметры метода --
        :param n_jobs: кол-во используемых рабочих процессов, при -1 берется число, равное CPU_LIMIT
        :param result_folder: папка, в которую будут сохраняться все результаты работы
        :param method: Метод автобиннинга: 'tree' - биннинг деревом, 'opt' - биннинг деревом с последующей оптимизацией границ бинов библиотекой optbinning
        :param validate: флаг для выполнения валидацонных тестов
        :param out: либо строка с названием эксель файла, либо объект pd.ExcelWriter для сохранения отчета
        :param save_model: название json файла для сохранения модели
        :param config: dict или json файл с параметрами. Должен содержать словарь вида:
                        {'DataSamples': {},       # параметры для инициализации объекта DataSamples, в случае если data имеет тип DataFrame
                         'WOE'        : {},       # параметры для инициализации объекта WOE
                         'auto_fit'   : {},       # параметры вызова метода WOE.auto_fit()
                         'mfa'        : {},       # параметры вызова метода self.mfa()
                        }
        """
        if isinstance(config, str):
            config = json.load(open(config, 'rt', encoding='utf-8'))
        elif not isinstance(config, dict):
            config = {}
        if data is not None:
            if isinstance(data, pd.DataFrame):
                ds_config = {'samples': {'Train': data},
                             'target': target,
                             'time_column': time_column,
                             'id_column': id_column,
                             'feature_descriptions': feature_descriptions,
                             'result_folder': result_folder,
                             'n_jobs': n_jobs,
                             'samples_split': {}}
                self.ds = DataSamples(**{**ds_config, **(config['DataSamples'] if 'DataSamples' in config else {})})
            else:
                self.ds = copy.copy(data)
                if n_jobs is not None:
                    self.ds.n_jobs = n_jobs
                if result_folder and isinstance(result_folder, str):
                    if not os.path.exists(result_folder):
                        os.makedirs(result_folder)
                    self.ds.result_folder = result_folder + ('/' if result_folder and not result_folder.endswith('/') else '')

        if self.transformer is None:
            woe_config = {'ds': self.ds}
            self.transformer = WOE(**{**woe_config, **(config['WOE'] if 'WOE' in config else {})})
            autofit_config = {'plot_flag': -1, 'method': method}
            self.transformer.auto_fit(**{**autofit_config, **(config['auto_fit'] if 'auto_fit' in config else {})})
            self.transformer.transform(self.ds)
        else:
            self.ds.logger.info('Using existing self.transformer')

        if not self.coefs:
            mfa_config = {'drop_corr_iteratively': False, 'result_file': out}
            self.mfa(**{**mfa_config, **(config['mfa'] if 'mfa' in config else {})})
        else:
            self.ds.logger.info('Using existing self.coefs')
            report_config = {'out': out}
            self.report(**{**report_config, **(config['report'] if 'report' in config else {})})

        if save_model:
            self.save_model(file_name=f'{self.ds.result_folder}{save_model}')
        if validate:
            self.validate(result_file='auto_validation.xlsx')

    def mfa(self, ds=None, features=None, hold=None, features_ini=None, limit_to_add=100, gini_threshold=10, psi_threshold=0.25,
            corr_method='pearson', corr_threshold=0.75, drop_with_most_correlations=False, drop_corr_iteratively=True,
            selection_type='stepwise', pvalue_threshold=0.05, pvalue_priority=False, scoring='gini', score_delta=0.1,
            n_stops=1, cv=None, drop_positive_coefs=True, crosses_simple=False, crosses_max_num=5,
            verbose=True, result_file='mfa.xlsx', plotbins_config=None, metrics=None, metrics_cv=None):
        """
        Многофакторный отбор. Проходит в 4 основных этапа:

        1) Отбор по джини и PSI. Исключаются все факторы с джини ниже gini_threshold и PSI выше psi_threshold.
            При заданном time_column PSI считается между временными срезами, иначе - между сэмплами

        2) Корреляционный анализ. Доступны два варианта работы:
            drop_with_most_correlations=False - итерационно исключается фактор с наименьшим джини из списка коррелирующих факторов
            drop_with_most_correlations=True - итерационно исключается фактор с наибольшим кол-вом коррелирующих с ним факторов

        3) Итерационный отобор. Доступны три типа отбора:
            selection_type='forward' - все доступные факторы помещаются в список кандидатов, на каждом шаге из списка кандидатов определяется лучший* фактор и добавляется в модель
            selection_type='backward' - в модель включаются все доступные факторы, затем на каждом шаге исключается худший* фактор
            selection_type='stepwise' - комбинация 'forward' и 'backward'. Каждый шаг состоит из двух этапов:
                    на первом из списка кандидатов отбирается лучший* фактор в модель,
                    на втором из уже включенных факторов выбирается худший* и исключается

            *Определение лучшего фактора:
            При pvalue_priority=False лучшим фактором считается тот, который увеличивает метрику scoring модели на наибольшую величину.
                Если величина такого увеличения ниже score_delta, то счетчик n_stops уменьшается на 1. Когда он достигнет нуля, отбор прекращается
            При pvalue_priority=True лучшим фактором считается фактор, который после добавления в модель имеет наименьшее p-value.
                Если величина этого p-value выше pvalue_threshold, то счетчик n_stops уменьшается на 1. Когда он достигнет нуля, отбор прекращается

            *Определение худшего фактора:
            Худшим фактором в модели считается фактор с наибольшим p-value.
                Если величина этого p-value ниже pvalue_threshold, то худший фактора не определяется, и исключения не происходит

        4) Если выставлен флаг drop_positive_coefs=True, то из списка отобранных на этапе 3 факторов итерационно
            исключаются факторы с положительными коэффициентами и факторы с p_value > pvalue_threshold
        :param ds: ДатаСэмпл. В случае, если он не содержит трансформированные переменные, то выполняется трансформация трансформером self.transformer.
                   При None берется self.ds
        :param features: исходный список переменных для МФА. При None берутся все переменные, по которым есть активный биннинг
        :param hold: список переменных, которые обязательно должны войти в модель
        :param features_ini: список переменных, с которых стартует процедура отбора. Они могут быть исключены в процессе отбора
        :param limit_to_add: максимальное кол-во переменных, которые могут быть добавлены к модели. Применимо только к типам forward и stepwise
        :param gini_threshold: граница по джини для этапа 1
        :param psi_threshold: граница по PSI для этапа 1
        :param corr_method: метод расчета корреляций для этапа 2. Доступны варианты 'pearson', 'kendall', 'spearman'
        :param corr_threshold: граница по коэффициенту корреляции для этапа 2
        :param drop_with_most_correlations: вариант исключения факторов в корреляционном анализе для этапа 2
        :param drop_corr_iteratively: исключение коррелирующих факторов не на отдельном этапе 2, а итеративно в процессе этапа 3
                                     (список кандидатов на добавление в модель формируется динамически после каждого шага,
                                      из него исключаются все коррелирующие с уже включенными факторы).
                                     Применимо только для типов отбора forward и stepwise
        :param selection_type: тип отбора для этапа 3
        :param pvalue_threshold: граница по p-value для этапа 3
        :param pvalue_priority: вариант определения лучшего фактора для этапа 3
        :param scoring: максимизируемая метрика для этапа 3.
                Варианты значений: 'gini', 'AIC', 'BIC' + все метрики доступные для вычисления через sklearn.model_selection.cross_val_score.
                Все информационные метрики после вычисления умножаются на -1 для сохранения логики максимизации метрики.
        :param score_delta: минимальный прирост метрики для этапа 3
        :param n_stops: количество срабатываний нарушений правил отбора по приросту метрики/p-value до завершения этапа 3
        :param cv: параметр cv для вызова sklearn.model_selection.cross_val_score. При None берется StratifiedKFold(5, shuffle=True)
        :param drop_positive_coefs: флаг для выполнения этапа 4

        --- Кросс переменные ---
        :param crosses_simple: True  - после трансформации кросс-переменные учавствут в отборе наравне со всеми переменными
                               False - сначала выполняется отбор только на основных переменных,
                                       затем в модель добавляются по тем же правилам кросс переменные, но не более, чем crosses_max_num штук
        :param crosses_max_num: максимальное кол-во кросс переменных в модели. учитывается только при crosses_simple=False

        --- Отчет ---
        :param verbose: флаг для вывода подробных комментариев в процессе работы
        :param result_file: название файла отчета. При None результаты не сохраняются
        :param plotbins_config: словарь с параметрами для построения графиков методом self.transformer.plot_bins()
        :param metrics: список метрик/тестов, результы расчета которых должны быть включены в отчет.
                          Элементы списка могут иметь значения (не чувствительно к регистру):
                              'ontime'  : расчет динамики джини по срезам,
                              'psi'     : расчет Population Population Stability Index,
                              metric:   : расчет любой доступной метрики методом DataSamples.calc_metric(metric=metric)
                              func      : пользовательская функция, которая принимает целевую и зависимую переменную,
                                          и возвращает числовое значение метрики

                                        Например,
                                        def custom_metric(y_true, y_pred):
                                            from sklearn.metrics import roc_curve, f1_score
                                            fpr, tpr, thresholds = roc_curve(y_true, y_pred)
                                            thres = thresholds[np.argmax(tpr * (1 - fpr))]
                                            return f1_score(y_true, (y_pred > thres).astype(int))
                                        metrics = ['vif', 'ks', 'psi', custom_metric]
        :param metrics_cv: список метрик, рассчитываемых через sklearn.model_selection.cross_val_score.
                          Аналогично параметру metrics элементами могут быть строки, поддерживаемые cross_val_score, либо пользовательские функции
                          Например, ['roc_auc', 'neg_log_loss', 'gini', 'f1', 'accuracy', custom_metric]

        """
        if ds is not None:
            self.ds = copy.copy(ds)
            if self.transformer is not None:
                self.ds = self.transformer.transform(self.ds, features=features, verbose=verbose)

        if features is None or self.transformer is not None:
            features = self.ds.features.copy()
        if not crosses_simple:
            cross_features = [f for f in features if is_cross_name(f)]
            features = [f for f in features if not is_cross_name(f)]
        else:
            cross_features = []
        if cv is None:
            cv = StratifiedKFold(5, shuffle=True, random_state=self.ds.random_state)
        elif self.ds is None:
            print('ERROR! No DataSamples to process')
            return None
        if self.ds.logger.level > 20:
            verbose = False
        if verbose:
            self.ds.logger.info(make_header('MFA', 150))
        if result_file and self.name and self.name not in result_file:
            result_file = self.name + '_' + result_file

        hold = set() if hold is None else {add_suffix(f) if add_suffix(f) in self.ds.features else f for f in hold}
        features_ini = hold if features_ini is None else {add_suffix(f) if add_suffix(f) in self.ds.features else f for f in features_ini} | hold
        features = list(set(features) | features_ini)

        gini_df = self.ds.calc_gini(add_description=True, features=features)
        ginis = gini_df[self.ds.train_name].to_dict()
        features = sorted(features, key=ginis.get(0), reverse=True)
        drop_features_gini = [f for f in features if abs(ginis.get(f, 0)) < gini_threshold and f not in features_ini]
        features = [f for f in features if f not in drop_features_gini]
        if verbose:
            self.ds.logger.info(make_header('Step 1', 100))
            self.ds.logger.info(f'Dropped {len(drop_features_gini)} features with gini lower {gini_threshold}: {drop_features_gini}')
        if psi_threshold < 1 and features:
            psi_df = self.ds.psi(features=features, time_column=self.ds.time_column)
            drop_features_psi = psi_df[(psi_df.max(axis=1) > psi_threshold) & (~psi_df.index.isin(features_ini))].index.to_list()
            features = [f for f in features if f not in drop_features_psi]
            dropped_df = gini_df[gini_df.index.isin(drop_features_gini + drop_features_psi)]\
                .merge(psi_df.set_axis([f'PSI {c}' for c in psi_df.columns], axis=1), how='left', left_index=True, right_index=True)
            if verbose:
                self.ds.logger.info(f'Dropped {len(drop_features_psi)} features with PSI higher {psi_threshold}: {drop_features_psi}')
        else:
            dropped_df = gini_df[gini_df.index.isin(drop_features_gini)]

        if not features:
            self.ds.logger.warning('Features set is empty. Break.')
            self.features = []
            return None
        if verbose:
            self.ds.logger.info(make_header('Step 2', 100))
            self.ds.logger.info(f'Performing correlation analysis for {len(features)} features...')
        if not drop_corr_iteratively:
            drop_features_corr = self.ds.CorrelationAnalyzer(sample_name=self.ds.train_name, features=features, hold=features_ini,
                                                             scores=ginis, drop_with_most_correlations=drop_with_most_correlations,
                                                             method=corr_method, threshold=corr_threshold, verbose=verbose)
            features = [f for f in features if f not in drop_features_corr]
            corr_df = None
        else:
            drop_features_corr = {}
            corr_df = self.ds.samples[self.ds.train_name][features].corr(method=corr_method).abs() > corr_threshold
            if verbose:
                self.ds.logger.info('Parameter drop_iteratively=True is selected, exclusion of correlated features will be performed during step 3 iteratively')
        if verbose:
            self.ds.logger.info(make_header('Step 3', 100))
        if selection_type in ['stepwise', 'forward', 'backward']:
            features, selection_fig = self.stepwise_selection(features=features, hold=hold, features_ini=features_ini,
                                                              limit_to_add=limit_to_add, verbose=verbose,
                                                              score_delta=score_delta, scoring=scoring, cv=cv,
                                                              pvalue_threshold=pvalue_threshold,
                                                              pvalue_priority=pvalue_priority, n_stops=n_stops,
                                                              selection_type=selection_type, corr_df=corr_df,
                                                              drop_positive_coefs=drop_positive_coefs)
        else:
            if selection_type != 'regularized':
                self.ds.logger.warning('Incorrect selection_type value. Set to regularized.')
            selection_fig = None
        if drop_positive_coefs and selection_type != 'stepwise':
            if verbose:
                self.ds.logger.info(make_header('Step 4', 100))
            features, regularized_fig = self.regularized_selection(features=features, hold=hold, scoring=scoring, cv=cv,
                                                                   pvalue_threshold=pvalue_threshold, verbose=verbose)
        else:
            regularized_fig = None
        if verbose:
            self.ds.logger.info(make_header('Final model', 100))
        self.fit(features=features)
        if result_file is not None and isinstance(result_file, (str, pd.ExcelWriter)):
            if isinstance(result_file, str):
                writer = pd.ExcelWriter(self.ds.result_folder + result_file,  engine='xlsxwriter')
            else:
                writer = result_file
            dropped_df.to_excel(writer, sheet_name='Dropped features')
            adjust_cell_width(writer.sheets['Dropped features'], gini_df)
            gini_df['Drop reason'] = gini_df.index.map(drop_features_corr)
            self.ds.corr_mat(sample_name=self.ds.train_name, features=[f for f in ginis if f not in dropped_df.index], description_df=gini_df, styler=True).to_excel(writer, sheet_name='Correlation analysis')
            adjust_cell_width(writer.sheets['Correlation analysis'], gini_df)
            self.report(out=writer, sheet_name=selection_type, pvalue_threshold=pvalue_threshold, verbose=verbose, add_figs=[selection_fig, regularized_fig],
                        gini_df=gini_df.drop(['Drop reason'], axis=1), metrics=metrics, metrics_cv=metrics_cv, cv=cv, plotbins_config=plotbins_config)
            if isinstance(result_file, str):
                writer.close()
                self.ds.logger.info(f'MFA report is saved to file {result_file}')

        if crosses_max_num > 0 and cross_features:
            if verbose:
                self.ds.logger.info('Processing cross features...')
            self.mfa(features=cross_features, hold=hold, gini_threshold=gini_threshold, psi_threshold=psi_threshold, corr_method=corr_method,
                    corr_threshold=corr_threshold, drop_with_most_correlations=drop_with_most_correlations,
                    drop_corr_iteratively=drop_corr_iteratively, selection_type='stepwise', pvalue_threshold=pvalue_threshold,
                    pvalue_priority=pvalue_priority, scoring=scoring, score_delta=score_delta, n_stops=n_stops, cv=cv,
                    drop_positive_coefs=drop_positive_coefs, features_ini=self.features,
                    limit_to_add=crosses_max_num, verbose=verbose, result_file=result_file[:-5] + '_cross.xlsx' if result_file else None,
                    crosses_simple=True)

    def report(self, ds=None, out='report.xlsx', sheet_name=None, pvalue_threshold=0.05, verbose=False, add_figs=None,
               gini_df=None, metrics=None, metrics_cv=None, cv=None, plotbins_config=None):
        """
        Генерация эксель отчета по обученной модели.
        :param ds: ДатаСэмпл. В случае, если он не содержит трансформированные переменные, то выполняется трансформация трансформером self.transformer
        :param out: либо строка с названием эксель файла, либо объект pd.ExcelWriter для сохранения отчета
        :param sheet_name: название листа в экселе
        :param pvalue_threshold: граница по p-value. Используется только для выделения значений p-value цветом
        :param verbose: флаг вывода комментариев в процессе работы
        :param add_figs: список из графиков, которые должны быть добавлены в отчет
        :param gini_df: датафрейм с джини всех переменных модели
        :param metrics: список метрик/тестов, результы расчета которых должны быть включены в отчет.
                          Элементы списка могут иметь значения (не чувствительно к регистру):
                              'ontime'  : расчет динамики джини по срезам,
                              'psi'     : расчет Population Population Stability Index,
                              metric:   : расчет любой доступной метрики методом DataSamples.calc_metric(metric=metric)
                              func      : пользовательская функция, которая принимает целевую и зависимую переменную,
                                          и возвращает числовое значение метрики

                                        Например,
                                        def custom_metric(y_true, y_pred):
                                            from sklearn.metrics import roc_curve, f1_score
                                            fpr, tpr, thresholds = roc_curve(y_true, y_pred)
                                            thres = thresholds[np.argmax(tpr * (1 - fpr))]
                                            return f1_score(y_true, (y_pred > thres).astype(int))
                                        metrics = ['vif', 'ks', 'psi', custom_metric]
        :param metrics_cv: список метрик, рассчитываемых через sklearn.model_selection.cross_val_score.
                          Аналогично параметру metrics элементами могут быть строки, поддерживаемые cross_val_score, либо пользовательские функции
                          Например, ['roc_auc', 'neg_log_loss', 'gini', 'f1', 'accuracy', custom_metric]
        :param cv: параметр cv для вызова sklearn.model_selection.cross_val_score
        :param plotbins_config: словарь с параметрами для построения графиков методом self.transformer.plot_bins()

        :return: ДатаФрейм со сводной таблицей отчета
        """
        if ds is None:
            ds = self.ds
        if sheet_name is None:
            sheet_name = self.name if self.name else 'model'
        if not self.coefs or not self.features:
            ds.logger.error('Please fit your model before calling this method.')
            return None
        if isinstance(out, str):
            writer = pd.ExcelWriter(add_ds_folder(ds, out), engine='xlsxwriter')
        elif isinstance(out, pd.ExcelWriter):
            writer = out
        else:
            ds.logger.error('Parameter out must have str or pd.ExcelWriter type.')
            return None
        if verbose:
            ds.logger.info('Generating report...')
        if metrics is None:
            metrics = ['wald', 'ks', 'vif', 'iv', 'psi', 'ontime']
        metrics = [t.lower() if isinstance(t, str) else t for t in metrics]
        score_field = 'model_score'
        gini_df = self.calc_gini(ds=ds, gini_df=gini_df)
        self.features = sorted(self.coefs, key=lambda x: gini_df[ds.train_name].to_dict().get(x, 0), reverse=True)
        self.coefs = {f: self.coefs[f] for f in self.features}
        if verbose:
            res_tbl = gini_df.replace(np.nan, '').rename_axis('').copy()
            res_tbl.iloc[:, 0] = res_tbl.iloc[:, 0].astype('str').str[:100]
            ds.logger.info('\n' + res_tbl.to_string())
        gini_df.columns = [('Gini', c) if c in (list(ds.samples) + ['CI_lower', 'CI_upper']) else (c, '') for c in gini_df.columns]
        for m in metrics:
            if m == 'wald':
                wald_df = pd.concat([self.wald_test(ds, sample_name=name, features=self.features)[['se', 'p-value'] if name == ds.train_name else ['p-value']].rename({'p-value': name}, axis=1) for name, sample in ds.samples.items()], axis=1)
                wald_df.columns = [('Wald p-value', c) if c in list(ds.samples) else (c, '') for c in wald_df.columns]
                gini_df = gini_df.merge(wald_df, left_index=True, right_index=True, how='left')
            elif isinstance(m, str) and m not in ['psi', 'ontime']:
                tmp = ds.calc_metric(metric=m, features=self.features + ([score_field] if m not in ['vif'] else []))
                if not tmp.empty:
                    tmp.columns = [(m, c) for c in tmp.columns]
                    gini_df = gini_df.merge(tmp, left_index=True, right_index=True, how='left')
            elif callable(m):
                gini_df = gini_df.merge(pd.concat([pd.DataFrame.from_dict({f: m(sample[ds.target], sample[f]) for f in self.features + [score_field]}, orient='index', columns=[(m.__name__, name)]) for name, sample in ds.samples.items()], axis=1), left_index=True, right_index=True, how='left')
        gini_df.columns = pd.MultiIndex.from_tuples(gini_df.columns)
        gini_df = gini_df.rename(index={score_field: 'model'})
        gini_df.style.applymap(lambda x: 'color: red' if x > pvalue_threshold else 'color: orange' if x > pvalue_threshold / 5 else 'color: black',
                               subset=pd.IndexSlice[:, [f for f in gini_df if f[0] == 'Wald p-value']]) \
            .to_excel(writer, sheet_name=sheet_name, startrow=1, startcol=0, float_format=f'%0.{self.round_digits}f')
        ds.corr_mat(features=self.features, styler=True).to_excel(writer, sheet_name=sheet_name, startrow=3, startcol=len(gini_df.columns) + 3)
        ws = writer.sheets[sheet_name]
        descr_len = len(ds.feature_descriptions.columns) if ds.feature_descriptions is not None else 0
        ws.set_column(0, 0 + descr_len, 30)
        ws.set_column(1 + descr_len, gini_df.shape[1], 15)
        ws.set_column(len(gini_df.columns) + 3, gini_df.shape[1] + 3, 30)
        ws.write(0, 0, 'Features in model:')
        ws.write(0, len(gini_df.columns) + 3, 'Correlations matrix:')
        m_col = 10
        if metrics_cv:
            model_metrics = []
            for m in metrics_cv:
                if callable(m):
                    try:
                        metric = {'Metric': m.__name__}
                        m = make_scorer(m)
                    except:
                        metric = {'Metric': str(m)}
                else:
                    metric = {'Metric': str(m)}
                for name, sample in ds.samples.items():
                    scores = cross_val_score(self.clf, sample[self.features], sample[ds.target], cv=cv, scoring=m if m != 'gini' else 'roc_auc')
                    if m == 'gini':
                        scores = np.array([2*x - 1 for x in scores])
                    metric[f'{name} mean'] = scores.mean()
                    metric[f'{name} std'] = scores.std()
                model_metrics.append(metric)
            ws.write(len(self.features) + 7, m_col, 'Model metrics on cross-validation:')
            pd.DataFrame(model_metrics).set_index('Metric').round(self.round_digits).to_excel(writer, sheet_name=sheet_name, startrow=len(self.features) + 8, startcol=m_col)
            m_col += len(ds.samples)*2 + 3
        fig_to_excel(self.roc_curve(ds, verbose=verbose, score_field=score_field), ws, row=0,
                     col=gini_df.shape[1] + max(len(self.features), len(ds.samples)) + 6)
        if add_figs:
            row = 0
            for fig in add_figs:
                if fig:
                    fig_to_excel(fig, ws, row=row, col=len(gini_df.columns) + max(len(self.features), len(ds.samples)) + 16)
                    row += 15

        if 'ontime' in metrics and ds.time_column:
            ws.write(len(self.features) + 7, m_col, 'Model Gini dynamics:')
            model_gini = ds.calc_gini_in_time(features=[score_field])*(-1)
            model_gini.columns = model_gini.columns.droplevel()
            model_gini.to_excel(writer, sheet_name=sheet_name, startrow=len(self.features) + 8, startcol=m_col)
            fig, ax = plt.subplots(1, 1, figsize=(8, 4))
            for name in ds.samples:
                ax.plot(model_gini.index, model_gini[name], label=name, marker='o')
            ax.fill_between(model_gini.index,
                            model_gini['CI_lower'],
                            model_gini['CI_upper'],
                            alpha=0.1, color='blue', label=f'{int(round((1 - ds.ci_alpha) * 100))}% CI')
            ax.set_ylabel('Gini')
            ax.set_title('Model Gini dynamics')
            for tick in ax.get_xticklabels():
                tick.set_rotation(30)
            ax.tick_params(axis='both', which='both', length=5, labelbottom=True)
            ax.xaxis.get_label().set_visible(False)
            ax.set_ylim(0, ax.get_ylim()[1])
            ax.legend(loc='center left', bbox_to_anchor=(1, 0.5), fontsize=12)
            fig_to_excel(fig, ws, row=len(self.features) + len(model_gini) + 11, col=m_col, scale=0.75)
            ws.write(len(self.features) + len(model_gini) + 27, m_col, 'Features Gini dynamics:')
            ds.calc_gini_in_time(features=self.features).to_excel(writer, sheet_name=sheet_name, startrow=len(self.features) + len(model_gini) + 28, startcol=m_col)

        if self.transformer is not None:
            ws.write(len(self.features) + 7, 0, 'Scorecard:')
            scorecard = self.transformer.export_scorecard(features=self.features, full=False)
            scorecard.set_index('feature').to_excel(writer, sheet_name=sheet_name,
                                                    startrow=len(self.features) + 8, startcol=0,
                                                    float_format=f'%0.{self.round_digits}f')
            figs = self.plot_bins(**{**{'plot_flag': verbose}, **(plotbins_config if plotbins_config else {})})
            for i, fig in enumerate(figs):
                fig_to_excel(fig, ws,
                             row=len(scorecard) + len(self.features) + 11 + i * (22 if ds.time_column is None else 32),
                             col=0, scale=0.7)
        else:
            scorecard = None

        if 'psi' in metrics:
            ds.psi(features=[score_field] + self.features, scorecard=scorecard, out=writer, sheet_name='PSI_samples')
            if ds.time_column:
                ds.psi(time_column=ds.time_column, features=[score_field] + self.features, scorecard=scorecard, out=writer, sheet_name='PSI_time')
        if isinstance(out, str):
            writer.close()
            ds.logger.info(f'The model report is saved to file {out}')
        plt.close('all')
        return gini_df

    def plot_bins(self, plot_flag=False, show_groups=False, verbose=False, all_samples=False, stat_size=None):
        return self.transformer.plot_bins(features=[rem_suffix(f) for f in self.features], folder=None, plot_flag=plot_flag,
                                          show_groups=show_groups, verbose=verbose, all_samples=all_samples, stat_size=stat_size)

    def calc_gini(self, ds=None, gini_df=None):
        """
        Вычисление джини модели
        :param ds: ДатаСэмпл
        :param gini_df: датафрейм с джини всех переменных модели. При None будет рассчитан автоматически

        :return: Датафрейм с джини
        """
        if ds is None:
            ds = self.ds
        ds = self.scoring(ds, score_field='model_score')
        if gini_df is None:
            gini_df = ds.calc_gini(add_description=True, features=self.features, abs=True)
        gini_df = gini_df.merge(pd.DataFrame.from_dict({**self.coefs, **{'intercept': self.intercept}}, orient='index',
                                   columns=['coefficient']), left_index=True, right_index=True, how='right')\
            .sort_values(by=ds.train_name, ascending=False, na_position='last')
        return pd.concat([gini_df, ds.calc_gini(features=['model_score'], abs=True)]).rename_axis('feature')

    def roc_curve(self, ds=None, score_field=None, verbose=True):
        """
        Рассчет джини модели на всех сэмплах и построение ROC-кривой
        :param ds: ДатаСэмпл
        :param score_field: поле со скором модели. При None рассчитывается автоматически
        :param verbose: флаг для вывода ROC-кривой в аутпут

        :return: ROC-кривая в виде plt.figure
        """
        if ds is None:
            ds = self.ds
        if score_field is None:
            score_field = 'tmp_score'
            ds = self.scoring(ds, score_field=score_field)
        fig = plt.figure(figsize=(4.5, 3))
        ax = fig.add_subplot(111)
        for name, sample in ds.samples.items():
            fpr, tpr, _ = roc_curve(sample[ds.target], sample[score_field])
            ax.plot(fpr, tpr, label=f'{name} (Gini = {round((auc(fpr, tpr) * 2 - 1) * 100, 2)})')
        ax.set_xlabel('False Positive Rate', fontsize=10)
        ax.set_ylabel('True Positive Rate', fontsize=10)
        ax.legend(fontsize=10)
        plt.tick_params(axis='both', which='both', labelsize=10)
        if verbose:
            plt.show()
        plt.close()
        return fig

    def validate(self, ds=None, result_file='validation.xlsx', score_field='score', pd_field='pd', scale_field=None):
        """
        Валидационные тесты модели на заданном ДатаСэмпле библиотекой bankmetrics
        :param ds: ДатаСэмпл
        :param result_file: результирующий эксель файл
        :param score_field: поле со расчитанным скором (при отсутcnвии этого поля в выборке будет вызван метод self.scoring)
        :param pd_field: поле с расчитанным PD
        :param scale_field: поле с расчитанным грейдом
        """
        if ds is None:
            ds = self.ds
        elif self.transformer is not None:
            ds = self.transformer.transform(ds)
        ds.logger.info(make_header('Validation', 150))

        if (score_field and score_field not in ds.samples[ds.train_name].columns) or \
                (pd_field and pd_field not in ds.samples[ds.train_name].columns) or \
                (scale_field and scale_field not in ds.samples[ds.train_name].columns):
            ds = self.scoring(ds,
                              score_field=score_field[:-7] if score_field and score_field.endswith('_calibr') else score_field,
                              pd_field=pd_field[:-7] if pd_field and pd_field.endswith('_calibr') else pd_field,
                              scale_field=scale_field[:-7] if scale_field and scale_field.endswith('_calibr') else scale_field)

        df = ds.samples[ds.train_name]
        with pd.ExcelWriter(ds.result_folder + result_file, engine='openpyxl') as writer:
            test_to_excel(Iv_test(df, col=self.features, cel=ds.target), writer.book)
            test_to_excel(Ks_test(df, s0=score_field, cel=ds.target), writer.book)
            test_to_excel(Gini_test(df, col=self.features, cel=ds.target, unidir=False), writer.book)
            test_to_excel(Cal_test(df, pd0=pd_field, cel=ds.target), writer.book)
            test_to_excel(Corr_test(df, col=self.features), writer.book)
            test_to_excel(Vif_test(df, col=self.features, s0=''), writer.book)
            test_to_excel(Woe_test(df, col=self.features, cel=ds.target), writer.book, test_name='WOE')
            if scale_field:
                test_to_excel(Bin_test(df, m0=scale_field, pd0=pd_field, cel=ds.target), writer.book)
                test_to_excel(Chi2_test(df, odr=ds.target, m=scale_field, pd2=pd_field), writer.book)
                test_to_excel(Herfindal_test(df, m0=scale_field), writer.book)
            if len(ds.samples) > 1:
                test_name = [name for name in ds.samples if name != ds.train_name][0]
                test_to_excel(Gini_model_test(df2=df, df1=ds.samples[test_name], name2=ds.train_name, name1=test_name, col=self.features, cel=ds.target, unidir=False), writer.book, test_name='Gini_model')
                if scale_field:
                    test_to_excel(Psi_test(df20=df, df1=ds.samples[test_name], name2=ds.train_name, name1=test_name, col=self.features, m0=scale_field), writer.book)
        ds.logger.info(f'Results of validation tests was saved to file {ds.result_folder + result_file}')
        plt.close('all')

    def set_scale(self, scale='master', pd_field=None, n=10):
        """
        Устанавливает шкалу для мэппинга PD
        :param scale: Шкала. Может задаваться в виде словаря {грейд: верхняя граница PD грейда}, либо принимать следующие значения:
                        'master': устанавливается мастер-шкала
                        'log'   : устанавливается логарифмическая шкала
                        'qcut'  : устанавливается шкала, разбивающая PD на n квантилей
                        'bins'  : устанавливается шкала с границами грейдов, рассчитанными биннингом PD
        :param pd_field: поле с PD, при None PD вычисляется моделью атоматически
        :param n: кол-во бинов для шкал 'log', 'qcut' и 'bins'
        """
        if isinstance(scale, dict):
            self.scale = scale
        elif scale == 'master':
            self.scale = {
                        'MA1': 0.0005,
                        'MA2': 0.000695,
                        'MA3': 0.000976,
                        'MB1': 0.001372,
                        'MB2': 0.001927,
                        'MB3': 0.002708,
                        'MC1': 0.003804,
                        'MC2': 0.005344,
                        'MC3': 0.007508,
                        'MD1': 0.010549,
                        'MD2': 0.014821,
                        'MD3': 0.020822,
                        'ME1': 0.029254,
                        'ME2': 0.041101,
                        'ME3': 0.057744,
                        'MF1': 0.081128,
                        'MF2': 0.11398,
                        'MF3': 0.160137,
                        'MG1': 0.224984,
                        'MG2': 0.31609,
                        'MG3': 1
                       }
        else:
            if not pd_field:
                self.scoring(score_field='tmp_score', pd_field='tmp_pd')
                pd_field = 'tmp_pd'
            if scale == 'log':
                splits = [np.exp(x) for x in  pd.qcut(np.log(self.ds.samples[self.ds.train_name][pd_field]), n, retbins=True, duplicates='drop')[1]]
            elif scale == 'qcut':
                splits = pd.qcut(self.ds.samples[self.ds.train_name][pd_field], n, retbins=True, duplicates='drop')[1]
            elif scale == 'bins':
                optb = optbinning.OptimalBinning(dtype='numerical', prebinning_method='cart',
                                                 max_n_prebins=n, min_prebin_size=0.01, monotonic_trend='auto_asc_desc')
                optb.fit(self.ds.samples[self.ds.train_name][pd_field], self.ds.samples[self.ds.train_name][self.ds.target])
                splits = optb.splits
            else:
                self.ds.logger.error("Incorrect value of scale! Must be a dictionary, or one of the values 'master', 'qcut', 'bins.")
                return
            self.scale = {f'({round(splits[i - 1], 3) if i > 1 else "-inf"}; {round(splits[i], 3) if i != len(splits) - 1 else "+inf"}]':
                              round(splits[i], 3) if i != len(splits) - 1 else np.inf
                          for i in range(1, len(splits))}
        self.ds.logger.info(f'Set scale {self.scale}')

    def to_scale(self, PD):
        if self.scale is None:
            return np.nan
        for s in self.scale:
            if PD < self.scale[s]:
                return s
        return 'MSD'

    def calibration_test(self, df, target, pd_field, w=1):
        """
        Тест на калибровку
        :param df: ДатаФрейм
        :param target: поле с таргетом
        :param pd_field: поле с рассчитанным PD
        :param w: вес дефолтных наблюдений для теста

        :return: ДатаФрейм с результатом теста
        """

        def group_test(df, label):
            ci_95 = sts.norm.interval(0.95)[1]
            ci_99 = sts.norm.interval(0.99)[1]
            n = df.shape[0]
            d = df[target].sum()
            nw = n - d + d * w
            pd1_v = df[pd_field].sum() / n
            odr_v = d * w / nw
            k = (np.abs(pd1_v * (1 - pd1_v) / n)) ** 0.5
            return [label, pd1_v, odr_v, n, d, nw, d * w,
                    pd1_v - k * ci_99, pd1_v - k * ci_95,
                    pd1_v + k * ci_95, pd1_v + k * ci_99,
                    '' if n == 0 else
                    'З' if np.abs(odr_v - pd1_v) < k * ci_95 else
                    'К' if np.abs(odr_v - pd1_v) > k * ci_99 else 'Ж']

        df['scale_calibr'] = df[pd_field].apply(self.to_scale)
        rows = []
        for group in self.scale:
            rows.append(group_test(df[df['scale_calibr'] == group], group))
        rows.append(group_test(df, 'All'))
        res = pd.DataFrame.from_dict(self.scale, orient='index', columns=['upper PD']) \
            .merge(pd.DataFrame(rows, columns=['grade', 'Avg PD', 'Target rate', 'Amount', 'Target amount',
                                               'Amount weighted', 'Target amount weighted',
                                               '1%', '5%', '95%', '99%', 'Signal']).set_index('grade'),
                   left_index=True, right_index=True, how='right')
        res.loc[(res['Avg PD'] > res['Target rate']) & (res['Signal'] != 'З'), 'Signal'] += '*'
        return res

    def calibrate(self, CT, ds=None, method=0, sample_name=None, score_field=None, result_file='calibration.xlsx',
                  plot_flag=True, fun=None, x0=None, args=None, lambda_ab=None):
        """
        Калибровка скора модели линейной функцией score_calibr = a + b*score. Результат сохраняется в self.calibration в виде списка [a, b]
        :param CT: значение центральной тенденции, к которому приводится среднее PD модели
        :param ds: ДатаСэмпл. При None берется self.ds
        :param method: метод калибровки. Доступны два варианта:
                        0 - Строится логрег скора на таргет, коэффициент b приравнивается полученному коэффициенту при скоре,
                            коэффицент a затем подбирается солвером для попадания в ЦТ при фиксированном b
                        1 - Расчитываются веса наблюдений и строится логрег скора на таргет с весами, a и b приравниваются коэффициентам логрега
                        2 - Коэффициенты рассчитываются минимизацией заданной функции через вызов scipy.optimize.minimize(fun=fun, x0=x0, args=args, method='nelder-mead')
                        любое другое значение - перерасчет коэффициентов не происходит, проводится тест на коэффициентах из self.calibration
        :param sample_name: название сэмпла, на котором проводится калибровка
        :param scale: шкала, на которой будет проведен биноминальный тест. Задается в виде словаря {грейд: верхняя граница PD грейда}. По умолчанию берется мастер-шкала
        :param score_field: поле со скором. Если оно отсутвует в ДатаСэмпле, то будет вызван метод self.scoring
        :param result_file: название эксель файла, в который будут записаны результаты
        :param plot_flag: флаг для вывода графика теста на калибровку

        --- Метод калибровки 2 ---
        :param fun: пользовательская функция. Должна иметь вид
             def fun(params, **args):
                ...
                return result_to_minimize

             где params - список из подбираемых параметров. Например, [a], [b], [a, b]
                 result_to_minimize - результирующее значение, которое будет минимизироваться солвером
        :param x0: начальные значения параметров
        :param args: кортеж аргументов
        :param lambda_ab: функция для формирования списка [a, b] из результирующих параметров солвера. При None берется lambda x: x

        Примеры использования для калибровки с ограничением минимального PD модели значением minPD:

        Вариант 1) Минимизация функции двух переменных
            def CT_minPD_delta(params, score, CT, minPD):
                a, b = params
                pd = 1 / (1 + np.exp(-(a + score * b)))
                return (CT - pd.mean()) ** 2 + 10 * (minPD - pd.min())**2

            fun = CT_minPD_delta
            x0 = [0, 1]
            args = (ds.samples[ds.train_name]['score'], CT, minPD)
            lambda_ab = None

        Вариант 2) Минимизация функции одной переменной, вычисление коэффициента b через связь minPD и лучшего скора по выборке
            def CT_delta(params, score, CT, minPD, best_score):
                a = params
                b = (-log(1 / minPD - 1) - a) / best_score
                pd = 1 / (1 + np.exp(-(a + score * b)))
                return (CT - pd.mean()) ** 2

            best_score = ds.samples[ds.train_name]['score'].min()
            fun = CT_delta
            x0 = 0
            args = (ds.samples[ds.train_name]['score'], CT, minPD, best_score)
            lambda_ab = lambda x: (x[0], (-log(1 / minPD - 1) - x[0]) / best_score))

        :return: коэффициенты калибровки [a, b]
        """
        if ds is None:
            ds = self.ds
        if sample_name is None:
            sample_name = ds.train_name
        df = ds.samples[sample_name].copy()
        if not score_field:
            score_field = 'tmp_score'
            df = self.scoring(df, score_field=score_field, pd_field=None)

        def CT_meanPD_b_fix(params, score, CT, b):
            a = params
            pd = 1 / (1 + np.exp(-(a + score * b)))
            return (CT - pd.mean()) ** 2
        # расчет веса дефолтных наблюдений

        n = df.shape[0]
        d = df[ds.target].sum()
        w = CT * (n - d) / ((1 - CT) * d)
        lr = copy.deepcopy(self.clf)
        if method == 0:
            lr.fit(df[[score_field]], df[ds.target])
            b = lr.coef_[0][0]
            res = minimize(fun=CT_meanPD_b_fix, x0=0, args=(df[score_field], CT, b), method='nelder-mead')
            self.calibration = [round(res['x'][0], self.round_digits + 1), round(b, self.round_digits + 1)]
        elif method == 1:
            lr.fit(df[[score_field]], df[ds.target], sample_weight=np.where(df[ds.target] == 1, w, 1))
            self.calibration = [round(lr.intercept_[0], self.round_digits + 1), round(lr.coef_[0][0], self.round_digits + 1)]
        elif method == 2:
            res = minimize(fun=fun, x0=x0, args=args, method='nelder-mead')
            if lambda_ab is None:
                lambda_ab = lambda x: x
            a, b = lambda_ab(res['x'])
            self.calibration = [round(a, self.round_digits + 1), round(b, self.round_digits + 1)]
        else:
            ds.logger.warning('Incorrect method. Using for calibration existing coefficients in self.calibration')

        if result_file is not None and self.scale:
            if isinstance(result_file, str):
                writer = pd.ExcelWriter(add_ds_folder(ds, result_file), engine='xlsxwriter')
            elif isinstance(result_file, pd.ExcelWriter):
                writer = result_file
            else:
                return self.calibration
            # расчет откалиброванного скора и PD
            if isinstance(self.calibration, list):
                df['score_calibr'] = self.calibration[0] + df[score_field] * self.calibration[1]
                df['PD_calibr'] = 1 / (1 + np.exp(-df['score_calibr']))

            # тест на калибровку для каждого грейда и всей выборки целиком
            res = self.calibration_test(df, ds.target, 'PD_calibr', w)
            res.style.applymap(lambda x: 'color: black' if not isinstance(x, str) or not x else
                                        'color: red' if x[0] == 'К' else
                                        'color: orange' if x[0] == 'Ж' else
                                        'color: green' if x[0] == 'З' else
                                        'color: black').to_excel(writer, sheet_name=sample_name, float_format=f'%0.{self.round_digits + 1}f')

            if isinstance(self.calibration, list):
                gini = DataSamples.get_f_gini(df[ds.target], df[score_field], abs_flag=True)
                pd.DataFrame(
                    [['CT', CT],
                     ['a', self.calibration[0]], ['b', self.calibration[1]],
                     ['вес дефолтных наблюдений', w], ['вес недефолтных наблюдений', 1],
                     ['min PD', df['PD_calibr'].min()],
                     ['Gini', gini],
                     ],
                    columns=['Key', 'Value']).to_excel(writer, sheet_name=sample_name, startrow=len(self.scale) + 4,
                                                       startcol=3, index=False, float_format='%0.5f')

            fig, ax = plt.subplots(1, 1, figsize=(10, 6))
            res = res[res.index != 'All']
            ax.set_ylabel('Observations')
            ax.set_xticks(range(res.shape[0]))
            ax.set_xticklabels(res.index, rotation=30, ha="right")
            amnt = res['Amount'].sum()
            ax.bar(range(res.shape[0]), (res['Amount'] - res['Target amount']) / amnt, zorder=0, color='forestgreen', label='Class 0')
            ax.bar(range(res.shape[0]), res['Target amount'] / amnt,  bottom=(res['Amount'] - res['Target amount']) / amnt, zorder=0, color='indianred', label='Class 1')
            ax.annotate('Amount:', xy=(-0.5, 1), xycoords=('data', 'axes fraction'), xytext=(0, 60), textcoords='offset pixels', color='black', size=8, ha='right')
            for i in range(res.shape[0]):
                ax.annotate(str(res['Amount'][i]), xy=(i, 1), xycoords=('data', 'axes fraction'),
                              xytext=(0, 60), textcoords='offset pixels', color='black', size=8, ha='center')
            ax.annotate('Target amount:', xy=(-0.5, 1), xycoords=('data', 'axes fraction'),
                        xytext=(0, 40), textcoords='offset pixels', color='black', size=8, ha='right')
            for i in range(res.shape[0]):
                ax.annotate(str(round(res['Target amount'][i])), xy=(i, 1), xycoords=('data', 'axes fraction'),
                            xytext=(0, 40), textcoords='offset pixels', color='black', size=8, ha='center')
            ax.annotate('Signal:', xy=(-0.5, 1), xycoords=('data', 'axes fraction'),
                        xytext=(0, 20), textcoords='offset pixels', color='black', size=8, ha='right')
            for i, S in enumerate(res['Signal'].values):
                ax.annotate(S, xy=(i, 1), xycoords=('data', 'axes fraction'),
                            xytext=(0, 20), textcoords='offset pixels',
                            color='red' if not S or S[0] == 'К' else 'orange' if S[0] == 'Ж' else 'green',
                            size=8, ha='center')
            ax.grid(False)
            ax.grid(axis='y', zorder=1)
            ax2 = ax.twinx()
            ax2.set_ylabel('Target rate (log scale)')
            ax2.set_yscale('log')
            ax2.grid(False)
            ax2.plot(range(res.shape[0]), res['Target rate'], 'bo-', linewidth=2.0, zorder=4, label='Target rate', color='black')
            ax2.fill_between(range(res.shape[0]), res['5%'], res['95%'],  alpha=0.1,  color='blue', label='95% CI')
            ax2.plot(range(res.shape[0]), res['Avg PD'], 'bo-', linewidth=2.0, zorder=4, label='Avg PD', color='blue')

            h1, l1 = ax.get_legend_handles_labels()
            h2, l2 = ax2.get_legend_handles_labels()
            ax.legend(h1 + h2, l1 + l2, fontsize=10, loc='upper left')
            fig.tight_layout()
            fig_to_excel(fig, writer.sheets[sample_name], row=0, col=len(res.columns) + 3)
            if plot_flag:
                plt.show()
            plt.close('all')
            if isinstance(result_file, str):
                writer.close()
                self.ds.logger.info(f'Calibration test is saved to file {result_file}')
        return self.calibration

    def predict_proba(self, ds=None, sample_name=None):
        """
        Вычисление вероятности целевого события
        :param ds: ДатаСэмпл. При None берется self.ds
        :param sample_name: название сэмпла для вычисление. При None берется ds.train_sample

        :return: np.array вероятностей
        """
        if ds is None:
            ds = self.ds
        if sample_name is None:
            sample_name = ds.train_name
        return 1 / (1 + np.exp(-(self.intercept + np.dot(ds.samples[sample_name][list(self.coefs.keys())], list(self.coefs.values())))))

    def wald_test(self, ds=None, sample_name=None, clf=None, features=None, fit=False):
        """
        Тест Вальда. Вычисление стандартной ошибки, Wald Chi-Square и p-value для всех коэффициентов модели на заданном ДатаСэмпле
        :param ds: ДатаСэмпл. При None берется self.ds
        :param sample_name: название сэмпла. При None берется ds.train_sample
        :param clf: классификатор модели. При None берется self.clf
        :param features: список переменных. При None берется self.features
        :param fit: флаг для обучения модели заново на текущих данных

        :return: дафрейм с полями 'feature', 'coefficient', 'se', 'wald', 'p-value'
        """
        if ds is None:
            ds = self.ds
        if features is None:
            features = self.features
        if sample_name is None:
            sample_name = ds.train_name
        if clf is None:
            coefs_list = [self.intercept] + list(self.coefs.values())
            predProbs = np.matrix([[1 - x, x] for x in self.predict_proba(ds, sample_name=sample_name)])
        else:
            if fit:
                clf.fit(ds.samples[sample_name][features], ds.samples[sample_name][ds.target])
            coefs_list = [round(x, self.round_digits) for x in (clf.intercept_.tolist() + clf.coef_[0].tolist())]
            predProbs = np.matrix(clf.predict_proba(ds.samples[sample_name][features]))
        coefs_list_to_check = [c for i, c in enumerate(coefs_list) if i == 0 or c != 0]
        features_to_check = [f for i, f in enumerate(features) if coefs_list[i+1] != 0]
        try:
            X_design = np.hstack((np.ones(shape = (ds.samples[sample_name][features_to_check].shape[0],1)),
                                  ds.samples[sample_name][features_to_check]))
            V = np.multiply(predProbs[:,0], predProbs[:,1]).A1
            covLogit = np.linalg.inv(np.matrix(X_design.T * V) * X_design)
            bse = np.sqrt(np.diag(covLogit))
            wald = (coefs_list_to_check / bse) ** 2
            pvalue = sts.chi2.sf(wald, 1)
        except:
            bse = [0 for _ in range(len(features_to_check) + 1)]
            wald = [0 for _ in range(len(features_to_check) + 1)]
            pvalue = [1 for _ in range(len(features_to_check) + 1)]
        return pd.DataFrame({'feature': ['intercept'] + features, 'coefficient': coefs_list})\
            .merge(pd.DataFrame({'feature': ['intercept'] + features_to_check, 'se': bse, 'wald': wald, 'p-value': pvalue}),
                   on='feature', how='left').set_index('feature')

    def regularized_selection(self, ds=None, features=None, hold=None, scoring='gini', cv=None, pvalue_threshold=0.05, verbose=False):
        """
        Отбор факторов на основе регуляризации - строится модель на всех переменных, затем итерационно исключаются
        переменные с нулевыми или положительными коэффициентами и низкой значимостью
        :param ds: ДатаСэмпл. При None берется self.ds
        :param features: исходный список переменных. При None берется self.features
        :param hold: список/сет переменных, которые обязательно должны остаться после отбора
        :param scoring: расчитываемый скор модели
        :param pvalue_threshold: граница значимости по p-value
        :param verbose: флаг для вывода подробных комментариев в процессе работы

        :return: кортеж (итоговый список переменных, график со скором в процессе отбора в виде объекта plt.figure)
        """
        if ds is not None:
            self.ds = ds

        if features is None:
            if self.features:
                features = self.features.copy()
            else:
                features = self.ds.features.copy()
        features = [add_suffix(f) if add_suffix(f) in self.ds.features else f for f in features]
        if hold is None:
            hold = set()
        else:
            hold = {add_suffix(f) if add_suffix(f) in features else f for f in hold}

        # correctness check
        for f in features:
            if f not in self.ds.samples[self.ds.train_name].columns:
                self.ds.logger.error(f'No {f} in DataSample!')
                return None

        ginis = self.ds.calc_gini(features=features, samples=[self.ds.train_name])[self.ds.train_name].to_dict()
        self.clf.fit(self.ds.samples[self.ds.train_name][features], self.ds.samples[self.ds.train_name][self.ds.target])
        scores = [self.get_cv_score(scoring=scoring, cv=cv, features=features, fit=False)]
        features_change = ['Initial']
        to_exclude = True
        if verbose:
            self.ds.logger.info(f'Dropping features with positive coefs and high p-values...')
        while to_exclude:
            to_exclude = None
            positive_to_exclude = {f: c for f, c in zip(features, self.clf.coef_[0]) if c > 0 and f not in hold}
            if positive_to_exclude:
                features_to_exclude = {x: ginis[x] for x in positive_to_exclude}
                to_exclude = min(features_to_exclude, key=features_to_exclude.get)
                add_text = f'coef: {positive_to_exclude[to_exclude]}'
            else:
                wald = self.wald_test(clf=self.clf, features=features)
                feature_to_exclude_array = wald[(wald['p-value'] > pvalue_threshold) & (wald['p-value'] == wald['p-value'].max()) & (wald.index.isin(list(hold) + ['intercept']) == False)].index.values
                if feature_to_exclude_array:
                    to_exclude = feature_to_exclude_array[0]
                    add_text = f'p-value: {wald[wald.index == feature_to_exclude_array[0]]["p-value"].values[0]}'
            if to_exclude:
                features.remove(to_exclude)
                features_change.append(f'- {to_exclude}')
                self.clf.fit(self.ds.samples[self.ds.train_name][features], self.ds.samples[self.ds.train_name][self.ds.target])
                new_score = self.get_cv_score(scoring=scoring, cv=cv, features=features, fit=False)
                scores.append(new_score)
                if verbose:
                    self.ds.logger.info(f'To drop: {to_exclude}, {scoring}: {new_score}, {add_text}')

        features = [f for f, c in zip(features, self.clf.coef_[0]) if c != 0]
        if len(features_change) > 1:
            fig = plt.figure(figsize=(max(len(features_change)//2, 5), 3))
            plt.plot(np.arange(len(scores)), scores, 'bo-', linewidth=2.0)
            plt.xticks(np.arange(len(features_change)), features_change, rotation=30, ha='right', fontsize=10)
            plt.tick_params(axis='y', which='both', labelsize=10)
            plt.ylabel(scoring)
            plt.title('Dropping features with positive coefs and high p-values')
            fig.tight_layout()
            if verbose:
                plt.show()
        else:
            fig = None
            if verbose:
                self.ds.logger.info(f'Nothing to drop')
        return features, fig

    @staticmethod
    def add_feature_stat(df, clf, scoring, cv):
        features = list(df.columns)[1:]
        clf.fit(df[features], df.iloc[:, 0])
        if clf.coef_[0][0] == 0:
            return 0, 1, 0
        coefs_list = [round(x, 2) for x in ([clf.intercept_[0]] + [c for c in clf.coef_[0] if c != 0])]
        features_to_check = [f for f, c in zip(features, clf.coef_[0]) if c != 0]
        try:
            predProbs = np.matrix(clf.predict_proba(df[features]))
            # Design matrix -- add column of 1's at the beginning of your X_train matrix
            X_design = np.hstack((np.ones(shape = (df[features_to_check].shape[0],1)),
                                  df[features_to_check]))
            V=np.multiply(predProbs[:,0], predProbs[:,1]).A1
            covLogit = np.linalg.inv(np.matrix(X_design.T * V) * X_design)
            bse = np.sqrt(np.diag(covLogit))
            wald = (coefs_list / bse) ** 2
            pvalue = sts.chi2.sf(wald, 1)
            if scoring.upper() in ['AIC', 'BIC', 'SIC',  'SBIC']:
                intercept_crit = np.ones((df.shape[0], 1))
                features_crit = np.hstack((intercept_crit, df[features_to_check]))
                scores_crit = np.dot(features_crit, coefs_list)
                ll = np.sum((df.iloc[:, 0] * scores_crit - np.log(np.exp(scores_crit) + 1)))
                if scoring.upper() == 'AIC':
                    score = 2 * len(coefs_list) - 2 * ll
                else:
                    score = len(coefs_list) * np.log(df.shape[0]) - 2 * ll
                score = -round(score, 3)
            else:
                try:
                    score = cross_val_score(clf, df[features], df.iloc[:, 0], cv=cv, scoring=scoring if scoring != 'gini' else 'roc_auc').mean()
                except:
                    score = 0
                if scoring == 'gini':
                    score = abs(round((2 * score - 1)*100, 2))
        except:
            return 0, 1, 0
        return score, pvalue[1], coefs_list[1]

    def add_feature(self, features=None, candidates=None, scoring='gini', cv=None, corr_df=None, n_plot=0):
        """
        Выполняет расчет метрик кандидатов для добавления в модель
        :param features: список переменных в модели. При None берется self.features
        :param candidates: список кандидатов на добавление. При None берется self.ds.features
        :param scoring: максимизируемая метрика.
                Варианты значений: 'gini', 'AIC', 'BIC'+ все метрики доступные для вычисления через sklearn.model_selection.cross_val_score.
                Все информационные метрики после вычисления умножаются на -1 для сохранения логики максимизации метрики.
        :param cv: параметр cv для вычисления скора sklearn.model_selection.cross_val_score
        :param corr_df: матрциа корреляций с элементами False/True. Если задана, то из списка кандидатов исключаются факторы,
                        которые имеют значения True как минимум с одним из списка features
        :param n_plot: кол-во фичей с начала списка, для которых нужно построить графики биннингов

        :return: ДатаФрейм с полями 'feature', 'score', 'p-value', 'coefficient', в котором для каждого кандидата приводятся
                 метрика scoring модели и p-value кандидата, после построения модели с переменными features + [feature]
        """
        result_list = []
        features = set(self.features) if features is None else set(features)
        candidates = set(self.ds.features) - features if candidates is None else set(candidates) - features
        if cv is None:
            cv = StratifiedKFold(5, shuffle=True, random_state=self.ds.random_state)
        if corr_df is not None and features:
            candidates = candidates - set(corr_df[corr_df[list(features)].max(axis=1)].index)
        clf = copy.deepcopy(self.clf)
        if self.ds.n_jobs > 1:
            jobs = {}
            iterations = len(candidates)
            candidates_iter = iter(candidates)
            with futures.ProcessPoolExecutor(max_workers=self.ds.n_jobs) as pool:
                while iterations:
                    for feature in candidates_iter:
                        jobs[pool.submit(self.add_feature_stat, df=self.ds.samples[self.ds.train_name][[self.ds.target, feature] + list(features)],
                                         clf=clf, scoring=scoring, cv=cv)] = feature
                        if len(jobs) > self.ds.n_jobs * 2:
                            break
                    for job in futures.as_completed(jobs):
                        iterations -= 1
                        feature = jobs[job]
                        result_list.append([feature] + list(job.result()))
                        del jobs[job]
                        break
        else:
            for feature in candidates:
                tmp_features = list(features) + [feature]
                clf.fit(self.ds.samples[self.ds.train_name][tmp_features],
                        self.ds.samples[self.ds.train_name][self.ds.target])
                score = self.get_cv_score(scoring=scoring, cv=cv, features=tmp_features, fit=False)
                try:
                    wald = self.wald_test(features=tmp_features, clf=clf, fit=False)
                    pvalue, coef = wald[wald.index == feature][['p-value', 'coefficient']].values[0]
                except:
                    pvalue, coef = 1, 0
                result_list.append([feature, score, pvalue, coef])
        result = pd.DataFrame(result_list, columns=['feature', 'score', 'p-value', 'coefficient']).sort_values(by=['score'], ascending=False).reset_index(drop=True)
        if n_plot and self.transformer is not None:
            self.transformer.plot_bins(features=[rem_suffix(f) for f in result['feature'].to_list()[:n_plot]])
        return result

    def drop_feature(self, features=None, scoring='gini', cv=None, n_plot=0):
        """
        Выполняет расчет метрик кандидатов для удаления из модели в модель
        :param features: список переменных в модели. При None берется self.features
        :param scoring: максимизируемая метрика.
                Варианты значений: 'gini', 'AIC', 'BIC' + все метрики доступные для вычисления через sklearn.model_selection.cross_val_score.
                Все информационные метрики после вычисления умножаются на -1 для сохранения логики максимизации метрики.
        :param cv: параметр cv для вычисления скора sklearn.model_selection.cross_val_score
        :param n_plot: кол-во фичей с начала списка, для которых нужно построить графики биннингов

        :return: ДатаФрейм с полями 'feature', 'score', 'p-value', 'coefficient', в котором для каждого кандидата приводятся
                 метрика scoring модели с переменными set(features) - {feature} и p-value кандидата
        """
        features = set(self.features) if features is None else set(features)
        if cv is None:
            cv = StratifiedKFold(5, shuffle=True, random_state=self.ds.random_state)
        if self.ds.n_jobs_restr > 1 and len(features) > 3:
            scores = {}
            jobs = {}
            iterations = len(features)
            candidates_iter = iter(features)
            with futures.ProcessPoolExecutor(max_workers=self.ds.n_jobs) as pool:
                while iterations:
                    for f in candidates_iter:
                        jobs[pool.submit(self.get_cv_score, scoring=scoring, cv=cv, fit=True,
                                        features=list(features - {f}))] = f
                        if len(jobs) > self.ds.n_jobs * 2:
                            break
                    for job in futures.as_completed(jobs):
                        iterations -= 1
                        scores[jobs[job]] = job.result()
                        del jobs[job]
                        break
        else:
            scores = {f: self.get_cv_score(scoring=scoring, cv=cv, features=list(features - {f}), fit=True) for f in features}
        result = pd.DataFrame.from_dict(scores, orient='index').reset_index().set_axis(['feature', 'score'], axis=1).sort_values(by=['score'], ascending=False)\
              .merge(self.wald_test(clf=copy.deepcopy(self.clf), features=list(features), fit=True).reset_index()[['feature', 'p-value', 'coefficient']],
                     on=['feature'], how='left').reset_index(drop=True)
        if n_plot and self.transformer is not None:
            self.transformer.plot_bins(features=[rem_suffix(f) for f in result['feature'].to_list()[:n_plot]])
        return result

    def stepwise_selection(self, ds=None, verbose=False, selection_type='stepwise', features=None, hold=None,
                           features_ini=None, limit_to_add=100, score_delta=0.01, scoring='gini', cv=None,
                           pvalue_threshold=0.05, pvalue_priority=False, corr_df=None, n_stops=1,
                           drop_positive_coefs=True):
        """
        selection_type='forward' - все доступные факторы помещаются в список кандидатов, на каждом шаге из списка кандидатов определяется лучший* фактор и добавляется в модель
            selection_type='backward' - в модель включаются все доступные факторы, затем на каждом шаге исключается худший* фактор
            selection_type='stepwise' - комбинация 'forward' и 'backward'. Каждый шаг состоит из двух этапов:
                    на первом из списка кандидатов отбирается лучший* фактор в модель,
                    на втором из уже включенных факторов выбирается худший* и исключается

            *Определение лучшего фактора:
            При pvalue_priority=False лучшим фактором считается тот, который увеличивает метрику scoring модели на наибольшую величину.
                Если величина такого увеличения ниже score_delta, то счетчик n_stops уменьшается на 1. Когда он достигнет нуля, отбор прекращается
            При pvalue_priority=True лучшим фактором считается фактор, который после добавления в модель имеет наименьшее p-value.
                Если величина этого p-value выше pvalue_threshold, то счетчик n_stops уменьшается на 1. Когда он достигнет нуля, отбор прекращается

            *Определение худшего фактора:
            Худшим фактором в модели считается фактор с наибольшим p-value.
                Если величина этого p-value ниже pvalue_threshold, то худший фактора не определяется, и исключения не происходит

        :param ds: ДатаСэмпл. При None берется self.ds
        :param verbose: флаг для вывода подробных комментариев в процессе работы
        :param selection_type: тип отбора. Варианты 'forward', 'backward', 'stepwise'
        :param features: исходный список переменных. При None берется self.features
        :param hold: список переменных, которые обязательно должны остаться после отбора
        :param features_ini: список переменных, с которых стартует отбор. Они могут быть исключены в процессе отбора
        :param limit_to_add: максимальное кол-во переменных, которые могут быть добавлены к модели
        :param score_delta: минимальный прирост метрики
        :param scoring: максимизируемая метрика.
                Варианты значений: 'gini', 'AIC', 'BIC' + все метрики доступные для вычисления через sklearn.model_selection.cross_val_score.
                Все информационные метрики после вычисления умножаются на -1 для сохранения логики максимизации метрики.
        :param cv: параметр cv для вычисления скора sklearn.model_selection.cross_val_score
        :param pvalue_threshold: граница значимости по p-value
        :param pvalue_priority: вариант определения лучшего фактора
        :param corr_df: матрциа корреляций с элементами False/True. Если задана, то после каждого шага из списка кандидатов исключаются факторы,
                        которые имеют значения True как минимум с одним из уже отобранных факторов
        :param n_stops: количество срабатываний нарушений правил отбора по приросту метрики/p-value до завершения процедуры отбора
        :param drop_positive_coefs: исключать переменные с положительным коэффициентов

        :return: кортеж (итоговый список переменных, график со скором в процессе отбора в виде объекта plt.figure)
        """
        def get_add_feature(features, candidates, score, n_stops):
            df_scores = self.add_feature(features, candidates, scoring, cv, corr_df)
            if drop_positive_coefs:
                df_scores = df_scores[df_scores['coefficient'] < 0]
            df_scores = df_scores[df_scores['p-value'] <= pvalue_threshold].drop(['coefficient'], axis=1)
            if not df_scores.empty:
                if pvalue_priority:
                    result = tuple(df_scores[df_scores['p-value'] == df_scores['p-value'].min()].values[0]) + ('Min p-value', )
                else:
                    result = tuple(df_scores[df_scores['score'] == df_scores['score'].max()].values[0])
                    if result[1] - score < score_delta:
                        n_stops -= 1
                        result += (f'Max {scoring} increase {round(result[1] - score, 2)} < {score_delta}, n_stops = {n_stops}', )
                    else:
                        result += (f'Max {scoring} increase {round(result[1] - score, 2) if score != -1000000 else round(result[1], 2)} > {score_delta}', )
                if n_stops > 0:
                    return result + (n_stops, )
            return None, score, 1, '', 0

        def get_drop_feature(features, hold, score):
            df_scores = self.drop_feature(features, scoring, cv)
            df_scores = df_scores[~df_scores['feature'].isin(list(hold))]
            if drop_positive_coefs and not df_scores[df_scores['coefficient'] >= 0].empty:
                df_scores = df_scores[df_scores['coefficient'] >= 0].fillna(1)
                result = tuple(df_scores[df_scores['p-value'] == df_scores['p-value'].max()].values[0])
                return result[:-1] + (f'Positive coef {result[-1]}', )
            df_scores.drop(['coefficient'], axis=1, inplace=True)
            if not df_scores.empty:
                pvalue_max = df_scores['p-value'].max()
                score_max = df_scores['score'].max()
                if pvalue_max > pvalue_threshold:
                    return tuple(df_scores[df_scores['p-value'] == pvalue_max].values[0]) + (f'P-value > {pvalue_threshold}', )
                elif score_max > score:
                    return tuple(df_scores[(df_scores['score'] == score_max)].values[0]) + (f'{scoring} increasing by {score_max - score}', )
            return None, score, 1, ''

        if ds is not None:
            self.ds = ds
        hold = set() if hold is None else set(hold)
        features = set(self.ds.features) if features is None else set(features)
        features_ini = hold if features_ini is None else set(features_ini) | hold
        candidates = features - features_ini
        if selection_type != 'backward' or features_ini != hold:
            features = features_ini.copy()
        if not verbose:
            self.ds.logger.setLevel(max(40, self.ds.logger_level))
        if features:
            score = self.get_cv_score(scoring=scoring, cv=cv, features=list(features), fit=True)
            graph = [(score, 'initial')]
            self.ds.logger.info(f'Initial features: {list(features)}, {scoring} score {score}')
        else:
            score = -1000000
            graph = []
        for f in features | features_ini:
            if f not in self.ds.samples[self.ds.train_name].columns:
                self.ds.logger.error(f'No {f} in DataSample!')
                return None
        self.ds.logger.info(f'{selection_type.capitalize()} selection of {len(candidates)} features starts...')
        self.ds.logger.info(f" n |{'Feature': ^60}|{scoring.capitalize(): ^10}|{'P-value': ^10}| Operation |{'Reason': ^40}")
        self.ds.logger.info('_' * 140)

        if selection_type in ['forward', 'backward', 'stepwise']:
            feature_sets = [features.copy()]
            for i in range(len(candidates)):

                if selection_type in ['forward', 'stepwise']:
                    f, score, pvalue, reason, n_stops = get_add_feature(features, candidates, score, n_stops)
                    if f:
                        features.add(f)
                        graph.append((score, f'+ {f}'))
                        self.ds.logger.info(f"{len(features): ^3}|{f: ^60}|{score: ^10}|{round(pvalue, 6): ^10}|{'add': ^11}|{reason: ^40}")
                    else:
                        self.ds.logger.info('No significant features to add were found')

                if selection_type in ['backward', 'stepwise']:
                    while features:
                        f, score, pvalue, reason = get_drop_feature(features, hold, score)
                        if f:
                            features.remove(f)
                            graph.append((score, f'- {f}'))
                            self.ds.logger.info(f"{len(features): ^3}|{f: ^60}|{score: ^10}|{round(pvalue, 6): ^10}|{'drop': ^11}|{reason: ^40}")
                        else:
                            break
                gc.collect()
                if features in feature_sets:
                    break
                feature_sets.append(features.copy())
                if selection_type == 'backward':
                    break
                if len(features - features_ini) >= limit_to_add:
                    self.ds.logger.info(f'Reached the limit of the number of added features in the model. Selection stopped.')
                    break

            gc.collect()
            features = feature_sets[-1]
        else:
            self.ds.logger.error('Incorrect kind of selection. Please use backward, forward or stepwise.')
            self.ds.logger.setLevel(self.ds.logger_level)
            return None
        fig = plt.figure(figsize=(max(len(graph)//2, 5), 3))
        plt.plot(np.arange(len(graph)), [x[0] for x in graph], 'bo-', linewidth=2.0)
        plt.xticks(np.arange(len(graph)), [x[1] for x in graph], rotation=30, ha='right', fontsize=8)
        plt.tick_params(axis='y', which='both', labelsize=8)
        plt.ylabel(scoring, fontsize=10)
        plt.title(f'{selection_type.capitalize()} score changes', fontsize=10)
        fig.tight_layout()
        if verbose:
            plt.show()
        self.ds.logger.setLevel(self.ds.logger_level)
        gc.collect()
        return list(features), fig

    def tree_selection(self, ds=None, selection_type='forward', model_type='xgboost', result_file='tree_selection.xlsx',
                       plot_pdp=False, verbose=False):
        """
        Выполняет отбо факторов методом autobinary.AutoSelection. На отоборанных факторах строится модель на бустинге и логреге.
        Если self.transformer не задан, то для отобранных факторов дополнительно делается автобиннинг
        :param ds: ДатаСэмпл с нетрансформированными переменными. При None берется self.ds
        :param selection_type: тип отбора. Варианты 'forward', 'backward', 'deep_backward'
        :param model_type: тип модели для классификации. Варианты 'xgboost', 'lightboost', 'catboost'
        :param result_file: файл, в который будут сохраняться результаты
        :param plot_pdp: флаг для построение графиков PDP. Нужна библиотека pdpbox
        :param verbose: флаг вывода комментариев в процессе работы
        """
        from autobinary.auto_permutation import PermutationSelection
        from autobinary.auto_selection import AutoSelection
        from autobinary.auto_trees import AutoTrees
        from autobinary.base_pipe import base_pipe
        if ds is None:
            ds = self.ds
        if plot_pdp:
            try:
                from autobinary.auto_pdp import PlotPDP
            except Exception as e:
                ds.logger.error(e)
                plot_pdp = False

        if model_type == 'xgboost':
            import xgboost
            params = {'eta': 0.01,
                      'n_estimators': 500,
                      'subsample': 0.9,
                      'max_depth': 6,
                      'objective': 'binary:logistic',
                      'n_jobs': ds.n_jobs,
                      'random_state': ds.random_state,
                      'eval_metric': 'logloss'}
            clf = xgboost.XGBClassifier(**params)
            fit_params = {
                'early_stopping_rounds': 100,
                'eval_metric': ['logloss', 'aucpr', 'auc'],
                'verbose': 25}
        elif model_type == 'lightboost':
            import lightgbm
            params = {'learning_rate': 0.01,
                      'n_estimators': 500,
                      'subsample': 0.9,
                      'max_depth': 6,
                      'objective': 'binary',
                      'metric': 'binary_logloss',
                      'n_jobs': ds.n_jobs,
                      'random_state': ds.random_state,
                      'verbose': -1}
            clf = lightgbm.LGBMClassifier(**params)
            fit_params = {
                'early_stopping_rounds': 100,
                'eval_metric': ['logloss', 'auc'],
                'verbose': -1}
        elif model_type == 'catboost':
            import catboost
            params = {'learning_rate': 0.01,
                      'iterations': 500,
                      'subsample': 0.9,
                      'depth': 6,
                      'loss_function': 'Logloss',
                      'thread_count': ds.n_jobs,
                      'random_state': ds.random_state,
                      'verbose': 0}
            clf = catboost.CatBoostClassifier(**params)

            fit_params = {
                'use_best_model': True,
                'early_stopping_rounds': 200,
                'verbose': 50,
                'plot': False}
        else:
            ds.logger.error('Wrong model_type!')
            return None

        X_train = ds.samples[ds.train_name][ds.features]
        y_train = ds.samples[ds.train_name][ds.target]
        prep_pipe = base_pipe(
            num_columns=[f for f in ds.features if f not in ds.cat_columns],
            cat_columns=ds.cat_columns,
            kind='all')
        prep_pipe.fit(X_train, y_train)
        new_X_train = prep_pipe.transform(X_train)

        perm_imp = PermutationSelection(
            model_type=model_type,
            model_params=params,
            task_type='classification')
        fi, fi_rank, depth_features, rank_features = perm_imp.depth_analysis(new_X_train, y_train, list(new_X_train.columns), 5)
        # задаем стратегию проверки
        strat = StratifiedKFold(
            n_splits=5,
            shuffle=True,
            random_state=ds.random_state)

        selection = AutoSelection(
            base_pipe=base_pipe,
            num_columns=[f for f in depth_features if f not in ds.cat_columns],
            cat_columns=[f for f in depth_features if f in ds.cat_columns],
            main_fit_params=fit_params,
            main_estimator=clf,
            X_train=X_train[depth_features],
            y_train=y_train,
            main_metric='gini',
            model_type=model_type)
        if selection_type == 'forward':
            selection_res = selection.forward_selection(strat=strat)
        elif selection_type == 'backward':
            selection_res = selection.backward_selection(strat=strat, first_degradation=True)
        elif selection_type == 'deep_backward':
            selection_res = selection.deep_backward_selection(strat=strat, tol=0.001)
        else:
            ds.logger.error('Wrong selection_type!')
            return None
        features = selection_res['features_stack']

        model = AutoTrees(
            main_estimator=clf,
            main_fit_params=fit_params,
            main_prep_pipe=prep_pipe,
            main_features=features,
            X_train=X_train[features],
            y_train=y_train,
            main_metric='gini',
            model_type=model_type)

        model.model_fit_cv(strat=strat)

        if plot_pdp:
            clf.fit(new_X_train, y_train)
            pdp_plot = PlotPDP(model=clf, X=new_X_train[features], main_features=features)
            pdp_plot.create_feature_plot(save=True, path=ds.result_folder + 'pdp_ice_plots', frac_to_plot=0.05)

        with pd.ExcelWriter(ds.result_folder + result_file, engine='xlsxwriter') as writer:
            samples = {}
            for name, sample in ds.samples.items():
                samples[name] = prep_pipe.transform(sample[features])
                samples[name][ds.target] = sample[ds.target].reset_index(drop=True)
            ds_tmp = DataSamples(samples=samples, target=ds.target, features=features, cat_columns=[], n_jobs=1, logger=None, verbose=False)
            if ds.bootstrap_base is not None:
                bootstrap_base = prep_pipe.transform(ds.bootstrap_base[features])
                bootstrap_base[ds.target] = ds.bootstrap_base[ds.target].reset_index(drop=True)
                ds_tmp.bootstrap_base = bootstrap_base
                ds_tmp.bootstrap = ds.bootstrap
            gini_df = ds_tmp.calc_gini(add_description=True)

            gini_df.to_excel(writer, sheet_name=model_type, startrow=1, startcol=0)
            ws = writer.sheets[model_type]
            adjust_cell_width(ws, gini_df)
            ws.write(0, 0, 'Features in model:')
            ws.write(0, len(gini_df.columns) + 3, 'Metrics:')
            model.get_extra_scores().round(self.round_digits).\
                to_excel(writer, sheet_name=model_type, startrow=1, startcol=len(gini_df.columns) + 3, index=False)
            if plot_pdp:
                for i, feature in enumerate(features):
                    ws.insert_image(len(features) + 5 + i * 28, 0, f'{ds.result_folder}pdp_ice_plots/PDP_{feature}.png',
                                    {'x_scale': 0.75, 'y_scale': 0.75})
            fig_to_excel(model.get_rocauc_plots(), ws, row=0,  col=len(gini_df.columns) + 10)

            if self.transformer is None:
                self.transformer = WOE(ds=ds, features=features)
                self.transformer.auto_fit(plot_flag=False, SM_on=False, BL_allow_Vlogic_to_increase_gini=20, G_on=False)
            ds = self.transformer.transform(ds, features=features, verbose=verbose)
            self.fit(ds, features=ds.features)
            self.report(ds, out=writer, sheet_name='logreg', pvalue_threshold=0.05, verbose=verbose)

    def fit(self, ds=None, sample_name=None, features=None):
        """
        Обучение модели
        :param ds: ДатаСэмпл. При None берется self.ds
        :param sample_name: название сэмпла на котором проводится обучение. При None берется ds.train_sample
        :param features: список переменных. При None берется self.features
        """
        if ds is not None:
            self.ds = ds
        if features:
            self.features = features
        elif not self.features:
            self.features = self.ds.features.copy()
        if sample_name is None:
            sample_name = self.ds.train_name
        try:
            self.clf.fit(self.ds.samples[sample_name][self.features], self.ds.samples[sample_name][self.ds.target])
            self.intercept = round(self.clf.intercept_[0], self.round_digits)
            self.coefs = {f: round(c, self.round_digits) for f, c in zip(self.features, self.clf.coef_[0])}
            self.ds.logger.info(f'intercept = {self.intercept}')
            self.ds.logger.info(f'coefs = {self.coefs}')
        except Exception as e:
            self.intercept = None
            self.coefs = {}
            self.ds.logger.error(e)
            self.ds.logger.error('Fit failed!')

    def get_cv_score(self, features=None, ds=None, sample_name=None, clf=None, cv=None, scoring='gini', fit=True):
        """
        Вычисление кросс-валидацией скора модели на заданной выборке
        :param ds: ДатаСэмпл. При None берется self.ds
        :param sample_name: название сэмпла. При None берется ds.train_sample
        :param clf: классификатор модели. При None берется self.clf
        :param cv: параметр cv для вычисления скора sklearn.model_selection.cross_val_score
        :param scoring: рассчитываемый скор. Варианты значений: 'gini', 'AIC', 'BIC' + все метрики доступные для вычисления через sklearn.model_selection.cross_val_score
        :param features: список переменных. При None берется self.features
        :param fit: флаг для обучения модели заново на текущих данных

        :return: рассчитанный скор
        """
        if ds is None:
            ds = self.ds
        if features is None:
            if self.features:
                features = self.features
            else:
                features = ds.features.copy()
        if not features:
            return np.nan
        if sample_name is None:
            sample_name = ds.train_name
        if clf is None:
            clf = copy.deepcopy(self.clf)
        if cv is None:
            cv = StratifiedKFold(5, shuffle=True, random_state=ds.random_state)
        if scoring.upper() in ['AIC', 'BIC', 'SIC',  'SBIC']:
            if fit:
                clf.fit(ds.samples[sample_name][features], ds.samples[sample_name][ds.target])
            features_kept = [f for f, c in zip(features, clf.coef_[0]) if c != 0]
            weights_crit = [clf.intercept_[0]] + [c for c in clf.coef_[0] if c != 0]
            intercept_crit = np.ones((ds.samples[sample_name].shape[0], 1))
            features_crit = np.hstack((intercept_crit, ds.samples[sample_name][features_kept]))
            scores_crit = np.dot(features_crit, weights_crit)
            ll = np.sum(ds.samples[sample_name][ds.target] * scores_crit - np.log(np.exp(scores_crit) + 1))

            if scoring.upper() == 'AIC':
                score = 2 * len(weights_crit) - 2 * ll
            else:
                score = len(weights_crit) * np.log(ds.samples[sample_name].shape[0]) - 2 * ll
            score = -round(score, 3)
        else:
            try:
                score = cross_val_score(clf, ds.samples[sample_name][features], ds.samples[sample_name][ds.target], cv=cv,
                                        scoring=scoring if scoring != 'gini' else 'roc_auc').mean()
            except:
                score = 0
            if scoring == 'gini':
                score = abs(round((2 * score - 1)*100, 2))
        return score

    def scoring(self, data=None, score_field='score', pd_field='pd', scale_field=None):
        """
        Скоринг выборки.
        :param data: ДатаСэмпл или ДатаФрейм. Возвращается объект того же типа
        :param score_field: поле, в которое должен быть записан посчитанный скор
        :param pd_field: поле, в которое должен быть записан посчитанный PD
        :param scale_field: поле, в которое должен быть записан посчитанный грейд

        :return: ДатаСэмпл или ДатаФрейм с добавленными полями скоров, PD и грейда
        """
        
        def score_df(df):
            df[score_field] = self.intercept + np.dot(df[list(self.coefs.keys())], list(self.coefs.values()))
            if pd_field:
                df[pd_field] = 1 / (1 + np.exp(-df[score_field]))
                if scale_field and self.scale:
                    df[scale_field] = df[pd_field].apply(self.to_scale)
            if self.calibration:
                df[f'{score_field}_calibr'] = self.calibration[0] + self.calibration[1] * df[score_field]
                if pd_field:
                    df[f'{pd_field}_calibr'] = 1 / (1 + np.exp(-df[f'{score_field}_calibr']))
                    if scale_field and self.scale:
                        df[f'{scale_field}_calibr'] = df[f'{pd_field}_calibr'].apply(self.to_scale)
            return df

        if not self.coefs or self.intercept is None:
            self.print_log(f'No calculated intercept and coefs in self, use self.fit() before. Return None', 30)
            return None
        if data is not None:
            if isinstance(data, pd.DataFrame):
                ds = DataSamples(samples={'Train': data}, features=[], cat_columns=[], logger=None, verbose=False)
            else:
                ds = data
            if self.transformer is not None:
                ds = self.transformer.transform(ds, features=list(self.coefs.keys()))
        else:
            ds = self.ds
        for name in ds.samples:
            ds.samples[name] = score_df(ds.samples[name])
        if ds.bootstrap_base is not None:
            ds.bootstrap_base = score_df(ds.bootstrap_base)
        if isinstance(data, pd.DataFrame):
            return ds.to_df(sample_field=None)
        else:
            return ds

    def get_code(self, score_field='score', pd_field=None, scale_field=None, lang='py'):
        result = ''
        if self.transformer is not None:
            for f_WOE in self.features:
                f = rem_suffix(f_WOE)
                start = f"df['{f_WOE}'] = " if lang == 'py' else ',\n    CASE\n'
                end = '\n' if lang == 'py' else f" as {f_WOE.replace('&', '_x_')}"
                if f in self.transformer.feature_woes:
                    result += self.transformer.feature_woes[f].get_transform_func(start=start, lang=lang) + end
                elif is_cross_name(f):
                    f1, f2 = cross_split(f)
                    result += self.transformer.feature_crosses[f1].get_transform_func(f2, start=start, lang=lang) + end
        if lang == 'py':
            result += f"df[{score_field}] = {self.intercept} + {' + '.join([f'''({c}) * df['{f}']''' for f, c in self.coefs.items()])}\n"
            if pd_field:
                result += f"df[{pd_field}] = 1 / (1 + np.exp(-df[{score_field}]))\n"
                if scale_field:
                    result += f"df[{scale_field}] = df[{pd_field}].apply(to_scale)\n"
            if self.calibration:
                result += f"df[{score_field} + '_calibr'] = {self.calibration[0]} + {self.calibration[1]} * df[{score_field}]\n"
                if pd_field:
                    result += f"df[{pd_field} + '_calibr'] = 1 / (1 + np.exp(-df[{score_field} + '_calibr']))\n"
                    if scale_field:
                        result += f"df[{scale_field} + '_calibr'] = df[{pd_field} + '_calibr'].apply(to_scale)\n"
        return result

    def to_py(self, file_name='model.py', score_field='score', pd_field='pd', scale_field=None):
        """
        Генерация хардкода функции scoring
        :param file_name: название питоновского файла, куда должен быть сохранен код
        :param score_field: поле, в которое должен быть записан посчитанный скор
        :param pd_field:  поле, в которое должен быть записан посчитанный PD
        :param scale_field:  поле, в которое должен быть записан посчитанный грейд
        """
        result = "import pandas as pd\nimport numpy as np\n\n"
        if scale_field:
            result += f'''
def to_scale(PD):
    scale = {self.scale}
    for s in scale:
        if PD < scale[s]:
            return s
    return 'MSD'\n\n'''

        result += f'''
def scoring(df, score_field='score', pd_field='pd', scale_field=None):
    """
    Функция скоринга выборки
    Arguments:
        df: [pd.DataFrame] входной ДатаФрейм, должен содержать все нетрансформированные переменные модели
        score_field: [str] поле, в которое должен быть записан посчитанный скор
        pd_field: [str] поле, в которое должен быть записан посчитанный PD
        scale_field: [str] поле, в которое должен быть записан посчитанный грейд
    Returns:
        df: [pd.DataFrame] выходной ДатаФрейм с добавленными полями трансформированных переменных, скоров, PD и грейда
    """
    '''
        result += self.get_code(score_field='score_field' if score_field else None,
                                pd_field='pd_field' if pd_field else None,
                                scale_field='scale_field' if scale_field else None,
                                lang='py').replace('\n','\n    ')
        result += f"return df\n\n"
        result += f'''df = scoring(df, score_field={f"'{score_field}'" if score_field else None}, pd_field={f"'{pd_field}'" if pd_field else None}, scale_field={f"'{scale_field}'" if scale_field else None})'''
        if file_name:
            file_name = add_ds_folder(self.ds, file_name)
            with open(file_name, 'w', encoding='utf-8') as file:
                file.write(result)
            self.print_log(f'The model code for implementation saved to file {file_name}')
        else:
            print(result)

    def print_log(self, s, level=20):
        if self.ds is not None and self.ds.logger is not None:
            if level == 20:
                self.ds.logger.info(s)
            elif level == 30:
                self.ds.logger.warning(s)
            elif level == 40:
                self.ds.logger.error(s)

    def to_sql(self, file_name='model.sql', score_field='score', pd_field='pd', scale_field=None, other_fileds=None,
               input_table='<input table>'):
        """
        Генерация хардкода функции scoring
        :param file_name: название питоновского файла, куда должен быть сохранен код
        :param score_field: поле, в которое должен быть записан посчитанный скор
        :param pd_field:  поле, в которое должен быть записан посчитанный PD
        :param scale_field:  поле, в которое должен быть записан посчитанный грейд
        """
        if other_fileds is None:
            other_fileds = []
        features = []
        for f in self.features:
            f = rem_suffix(f)
            if is_cross_name(f):
                features += cross_split(f)
            else:
                features.append(f)
        result = f"SELECT t.*, {self.intercept} + {' + '.join([f'({c}) * {f}' for f, c in self.coefs.items()])} as {score_field}\nFROM (\n"
        result += f"    SELECT {', '.join(other_fileds + features).replace('&', '_x_')}"

        result += self.get_code(score_field='score_field' if score_field else None,
                                pd_field='pd_field' if pd_field else None,
                                scale_field='scale_field' if scale_field else None,
                                lang='sql').replace('\n', '\n    ')
        result += f"\n    FROM {input_table}\n    ) as t\n"
        if file_name:
            file_name = add_ds_folder(self.ds, file_name)
            with open(file_name, 'w', encoding='utf-8') as file:
                file.write(result)
            self.print_log(f'The model SQL code for implementation saved to file {file_name}')
        else:
            print(result)

    def etalons(self, ds=None, n_rows=10000, file_in='etalon_in.csv', file_out='etalon_out.csv', add_fields=None):
        """
        Создает входной и выходной эталоны для модели.
        :param ds: ДатаСэмпл на основе короторго формируются эталоны. При None берется self.ds
        :param n_rows: Кол-во строк в эталоне
        :param file_in: Файл для сохранения входного эталона
        :param file_out: Файл для сохранения выходного эталона
        :param add_fields: Список дополнительных полей для включения в эталон
        return Выходной эталон
        """
        if ds is None:
            ds = self.ds
        if add_fields is None:
            add_fields = []
        features = list({f: None for l in [cross_split(f) for f in self.features] for f in l if f})
        add_features = [f for f in [ds.target, ds.time_column, ds.id_column] if f is not None] + list(add_fields)
        df = ds.to_df()
        df = df.sample(min(n_rows, len(df)))[list(dict.fromkeys(add_features + features))]
        df.to_csv(file_in, index=False)
        df = self.scoring(df)
        df.to_csv(file_out, index=False)
        return df

    def to_mlflow(self, experiment_name, run_name, spark, data):
        """
        Сохраняет модель в MlFlow используя библиотеку vtb_mlops
        :param experiment_name: Строка с идентификатором эксперимента
        :param run_name: Строка с названием запуска
        :param spark: Авторизованная Spark сессия
        :param data: Список артефактов модели. Элемент - путь до файла
        """
        import getpass
        try:
            import vtb_mlops
        except:
            self.print_log('vtb_mlops not found', 40)
            self.print_log('Please use "pip install vtb_mlops" first', 20)
            return None
        os.environ['MLFLOW_TRACKING_USERNAME'] = input("Введите свой логин: ")
        os.environ['MLFLOW_TRACKING_PASSWORD'] = getpass.getpass("Введите свой пароль: ")
        os.environ['MLFLOW_TRACKING_URI'] = 'https://mlflow-smb.apps.ps5-lpap02.region.vtb.ru'
        self.etalons()
        vtb_mlops.save_mlflow(model=self, experiment_name=experiment_name, run_name=run_name, etalon_in_csv='etalon_in.csv', etalon_out_csv='etalon_out.csv', spark=spark, data=data)

    def report_doc(self, report_name='report.docx'):
        """
        Создает отчет о разработке в docx формате
        :param report_name: назание файла для сохранения отчета
        """
        from docx import Document
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from docx.shared import Inches, Pt, Mm
        from docx.enum.text import WD_BREAK
        doc = Document(pkg_resources.resource_filename(__name__, 'templates/template_1.docx'))
        score_field = 'model_score'
        self.ds = self.scoring(self.ds, score_field=score_field)
        scorecard = self.transformer.export_scorecard(features=self.features, full=False)[['feature', 'group', 'values', 'woe', 'special_bins']]
        tags = {'@str!model_name@': self.name if self.name else 'No name',
                '@str!year@': str(datetime.datetime.now().year),
                '@img!binnings@': self.transformer.plot_bins(features=[rem_suffix(f) for f in self.features],
                                                             folder=None, plot_flag=False),
                '@tbl!factors_table@': self.ds.feature_descriptions.reset_index() if self.ds.feature_descriptions is not None else pd.DataFrame(),
                '@tbl!scorecard@': scorecard,
                '@tbl!samples_stat@': pd.DataFrame([{'': name, 'Количество наблюдений': sample.shape[0],
                                                     'Количество ЦС': sample[self.ds.target].sum(),
                                                     'Доля ЦС': round(sample[self.ds.target].sum() / sample.shape[0],
                                                                      3),
                                                     'Период': f'{sample[self.ds.time_column].min()} - {sample[self.ds.time_column].max()}' if self.ds.time_column is not None else 'NA'}
                                                    for name, sample in self.ds.samples.items()]).set_index(
                    '').T.reset_index(),
                '@tbl!report_table@': self.calc_gini().drop(['CI_lower', 'CI_upper'], axis=1).reset_index().replace(np.nan, ''),
                '@case!start_calibration@': (self.calibration is not None, '@case!end_calibration@', "Калибровка модели не выполнялась.")
                }
        tag_re = re.compile('@\w*!\w*@')

        del_until = ''
        paragraphs_to_delete = []
        for paragraph in doc.paragraphs:
            if del_until:
                if del_until not in paragraph.text:
                    paragraphs_to_delete.append(paragraph)
                    continue
                else:
                    paragraph.text = paragraph.text.replace(del_until, "")
                    del_until = ''

            for tag in tag_re.findall(paragraph.text):
                if tag.startswith('@case!'):
                    if tag in tags and not tags[tag][0]:
                        del_until = tags[tag][1]
                        paragraph.text = paragraph.text.replace(tag, tags[tag][2])
                    else:
                        paragraph.text = paragraph.text.replace(tag, '')
                elif tag in tags:
                    if tag.startswith('@str!'):
                        paragraph.text = paragraph.text.replace(tag, tags[tag])
                    elif tag.startswith('@img!'):
                        paragraph.text = paragraph.text.replace(tag, '')
                        r = paragraph.add_run()
                        for fig in tags[tag]:
                            imgdata = io.BytesIO()
                            fig.savefig(imgdata, format='png', bbox_inches="tight")
                            r.add_picture(imgdata, width=Mm((doc.sections[0].page_width - doc.sections[0].left_margin -
                                                             doc.sections[0].right_margin) / 36000))


        for paragraph in paragraphs_to_delete:
            p = paragraph._element
            p.getparent().remove(p)

        for table in doc.tables:
            try:
                tag = tag_re.search(table.cell(1, 0).text)[0]
            except:
                continue
            if tag in tags and tag.startswith('@tbl!'):
                tbl_df = tags[tag]
                if len(tbl_df.columns) > len(table.rows[0].cells):
                    for i in range(len(table.rows[0].cells), len(tbl_df.columns)):
                        table.add_column(Inches(1.5))
                        table.cell(0, i).text = tbl_df.columns[i]
                for i, row in tbl_df.iterrows():
                    if i == 0:
                        row_cells = table.row_cells(1)
                    else:
                        row_cells = table.add_row().cells
                    for j, v in enumerate(row):
                        row_cells[j].text = str(v)
        report_name = add_ds_folder(self.ds, report_name)
        doc.save(report_name)
        self.ds.logger.info(f'Report saved to file {report_name}')

class NpEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, np.integer):
            return int(obj)
        if isinstance(obj, np.floating):
            return float(obj)
        if isinstance(obj, np.ndarray):
            return obj.tolist()
        return super(NpEncoder, self).default(obj)
